"""
周报助手 - 云端全自动版
用户只需扫码登录，其他全自动
"""

import asyncio
import io
import sys
import csv
import re
import base64
from pathlib import Path
from datetime import datetime, timedelta
from io import StringIO

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from fastapi import FastAPI
from fastapi.responses import HTMLResponse, Response, JSONResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配置
TEMPLATE = Path("template.docx")
DEFAULT_WPS_URL = "https://www.kdocs.cn/l/cmvjNWclJM5P"

app = FastAPI(title="周报助手")

# 任务存储
tasks = {}


def parse_table(text):
    """解析表格"""
    reader = csv.reader(StringIO(text), delimiter='\t')
    rows = []
    for row in reader:
        rows.append([cell.strip() for cell in row])
    return rows


def set_font(cell, font='仿宋', size=11, bold=False):
    """设置字体"""
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = font
            r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            r.font.size = Pt(size)
            r.font.bold = bold


def add_borders(table):
    """加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for n in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)
    if not tbl.tblPr:
        tbl.insert(0, tblPr)


def generate_report(table_content: str, date: str) -> bytes:
    """生成周报文档"""
    data = parse_table(table_content)
    if not data:
        raise ValueError("表格内容为空")
    
    doc = Document(str(TEMPLATE))
    
    # 删除旧表格
    for t in list(doc.tables):
        t._tbl.getparent().remove(t._tbl)
    
    # 修改标题日期
    for para in doc.paragraphs:
        full_text = ''.join([r.text for r in para.runs])
        if '工作周报' in full_text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = f'工作周报（{date}）'
            para.runs[0].font.name = '微软雅黑'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            para.runs[0].font.size = Pt(20)
            para.runs[0].font.bold = True
            break
    
    # 找详细工单位置
    detail = None
    for para in doc.paragraphs:
        if '详细工单' in para.text:
            detail = para
            break
    
    cols = max(len(r) for r in data)
    
    # 表格1
    t1 = doc.add_table(len(data), cols)
    add_borders(t1)
    for i, row in enumerate(data):
        for j, txt in enumerate(row):
            t1.rows[i].cells[j].text = txt
            set_font(t1.rows[i].cells[j], bold=(i==0))
    
    if detail:
        detail._element.addprevious(t1._tbl)
    
    # 表格2
    data2 = data + [['其他需汇报信息：无', '', '', '']]
    t2 = doc.add_table(len(data2), cols)
    add_borders(t2)
    for i, row in enumerate(data2):
        for j, txt in enumerate(row):
            t2.rows[i].cells[j].text = txt
            set_font(t2.rows[i].cells[j], bold=(i==0))
    
    if detail:
        detail._element.addnext(t2._tbl)
    
    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


async def run_browser_task(task_id: str, wps_url: str, input_date: str):
    """后台运行浏览器任务"""
    from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeout
    
    browser = None
    playwright = None
    
    try:
        tasks[task_id]["status"] = "launching"
        tasks[task_id]["message"] = "正在启动浏览器..."
        
        playwright = await async_playwright().start()
        
        # 启动浏览器（无头模式，但添加隐藏自动化特征的参数）
        browser = await playwright.chromium.launch(
            headless=True,
            args=[
                '--disable-blink-features=AutomationControlled',
                '--disable-infobars',
                '--no-sandbox',
                '--disable-dev-shm-usage',
            ]
        )
        
        context = await browser.new_context(
            viewport={'width': 1280, 'height': 800},
            user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        )
        page = await context.new_page()
        
        # 注入脚本隐藏自动化特征
        await page.add_init_script("""
            Object.defineProperty(navigator, 'webdriver', {get: () => undefined});
            window.chrome = {runtime: {}};
            Object.defineProperty(navigator, 'plugins', {get: () => [1, 2, 3, 4, 5]});
            Object.defineProperty(navigator, 'languages', {get: () => ['zh-CN', 'zh']});
        """)
        
        # 打开WPS
        tasks[task_id]["status"] = "loading"
        tasks[task_id]["message"] = "正在打开WPS表格..."
        
        try:
            await page.goto(wps_url, wait_until="domcontentloaded", timeout=30000)
        except:
            pass
        
        await asyncio.sleep(3)
        
        # 检测是否需要登录
        tasks[task_id]["status"] = "waiting_login"
        tasks[task_id]["message"] = "请在下方扫码登录微信"
        
        # 等待登录，每2秒截图一次
        max_wait = 90  # 最多等待90秒
        logged_in = False
        
        for i in range(max_wait // 2):
            try:
                # 截图
                screenshot = await page.screenshot(type="jpeg", quality=60)
                screenshot_b64 = base64.b64encode(screenshot).decode('utf-8')
                tasks[task_id]["screenshot"] = screenshot_b64
                tasks[task_id]["wait_time"] = i * 2
                
                # 检查URL
                current_url = page.url
                if 'passport' not in current_url and 'singlesign' not in current_url and 'login' not in current_url:
                    # 检查是否有表格元素
                    try:
                        canvas = await page.query_selector('[class*="canvas"], [class*="sheet"], [class*="editor"], canvas')
                        if canvas:
                            logged_in = True
                            break
                    except:
                        pass
                
                await asyncio.sleep(2)
            except Exception as e:
                await asyncio.sleep(2)
        
        if not logged_in:
            # 最后再检查一次
            current_url = page.url
            if 'passport' not in current_url and 'singlesign' not in current_url:
                logged_in = True
        
        if not logged_in:
            raise Exception("登录超时，请刷新页面重试")
        
        # 登录成功
        tasks[task_id]["status"] = "reading"
        tasks[task_id]["message"] = "登录成功，正在读取表格..."
        tasks[task_id]["screenshot"] = None
        
        # 等待表格加载
        await asyncio.sleep(5)
        
        # 点击最新Sheet
        sheet_name = ""
        try:
            # 尝试多种选择器
            selectors = [
                '[class*="sheet-tab"]:not([style*="display: none"])',
                '.sheet-tab:last-child',
                '[role="tab"]:last-child',
            ]
            for selector in selectors:
                try:
                    tabs = await page.query_selector_all(selector)
                    if tabs:
                        last_tab = tabs[-1]
                        sheet_name = await last_tab.inner_text()
                        sheet_name = sheet_name.strip()
                        await last_tab.click()
                        await asyncio.sleep(2)
                        break
                except:
                    continue
        except Exception as e:
            pass
        
        # 获取表格数据
        tasks[task_id]["message"] = "正在提取表格数据..."
        
        # 方法1：使用JavaScript提取表格
        table_text = await page.evaluate('''() => {
            // 尝试从剪贴板API获取
            // 或者遍历DOM获取表格内容
            
            let result = [];
            
            // 查找所有单元格
            const cells = document.querySelectorAll('td, th, [class*="cell"]');
            if (cells.length > 0) {
                let rows = {};
                cells.forEach(cell => {
                    const rect = cell.getBoundingClientRect();
                    const y = Math.round(rect.top / 25); // 按行分组
                    if (!rows[y]) rows[y] = [];
                    rows[y].push(cell.innerText?.trim() || '');
                });
                
                const sortedYs = Object.keys(rows).sort((a, b) => a - b);
                sortedYs.forEach(y => {
                    result.push(rows[y].join('\\t'));
                });
                
                return result.join('\\n');
            }
            
            // 备选：获取整个编辑区内容
            const editor = document.querySelector('[class*="editor"], [class*="canvas"], [contenteditable="true"]');
            if (editor) {
                return editor.innerText;
            }
            
            return '';
        }''')
        
        # 方法2：如果方法1失败，尝试全选复制
        if not table_text or len(table_text) < 20:
            try:
                # 点击编辑区域
                await page.click('[class*="canvas"], [class*="editor"], body', timeout=3000)
            except:
                pass
            await asyncio.sleep(0.5)
            
            # 全选
            await page.keyboard.press('Control+A')
            await asyncio.sleep(0.5)
            
            # 复制
            await page.keyboard.press('Control+C')
            await asyncio.sleep(1)
            
            # 使用JavaScript读取选中内容
            table_text = await page.evaluate('''() => {
                const selection = window.getSelection();
                if (selection.rangeCount > 0) {
                    const range = selection.getRangeAt(0);
                    const div = document.createElement('div');
                    div.appendChild(range.cloneContents());
                    return div.innerText || '';
                }
                return '';
            }''')
        
        # 关闭浏览器
        await browser.close()
        await playwright.stop()
        browser = None
        playwright = None
        
        # 确定日期
        date = input_date
        if not date:
            match = re.search(r'(\d{4}[-—–]\d{4})', sheet_name)
            if match:
                date = match.group(1).replace('—', '-').replace('–', '-')
            else:
                date = datetime.now().strftime("%m%d") + "-" + (datetime.now() + timedelta(days=4)).strftime("%m%d")
        
        tasks[task_id]["date"] = date
        tasks[task_id]["message"] = "正在生成周报..."
        
        # 生成报告
        if not table_text or len(table_text) < 10:
            raise ValueError("未能读取表格内容，请确保已登录并选择了正确的Sheet")
        
        doc_bytes = generate_report(table_text, date)
        tasks[task_id]["document"] = base64.b64encode(doc_bytes).decode('utf-8')
        tasks[task_id]["status"] = "done"
        tasks[task_id]["message"] = "周报已生成！"
        
    except Exception as e:
        tasks[task_id]["status"] = "error"
        tasks[task_id]["message"] = f"错误: {str(e)}"
    finally:
        if browser:
            try:
                await browser.close()
            except:
                pass
        if playwright:
            try:
                await playwright.stop()
            except:
                pass


# HTML页面
HTML_PAGE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>周报助手</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Microsoft YaHei', sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
        }
        .container {
            background: white;
            border-radius: 16px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            padding: 40px;
            max-width: 800px;
            width: 100%;
        }
        h1 { color: #333; margin-bottom: 10px; font-size: 28px; text-align: center; }
        .subtitle { color: #666; margin-bottom: 30px; font-size: 14px; text-align: center; }
        
        label { display: block; font-weight: 600; color: #333; margin-bottom: 8px; }
        input[type="text"] {
            width: 100%;
            padding: 12px 16px;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            font-size: 16px;
            margin-bottom: 20px;
        }
        input[type="text"]:focus { outline: none; border-color: #667eea; }
        
        .btn {
            width: 100%;
            padding: 14px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 8px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: transform 0.2s;
        }
        .btn:hover { transform: translateY(-2px); }
        .btn:disabled { background: #ccc; cursor: not-allowed; transform: none; }
        
        .msg {
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
            text-align: center;
        }
        .msg.info { background: #e3f2fd; color: #1565c0; }
        .msg.success { background: #e8f5e9; color: #2e7d32; }
        .msg.error { background: #ffebee; color: #c62828; }
        
        .screenshot-container {
            margin: 20px 0;
            text-align: center;
            display: none;
        }
        .screenshot-container img {
            max-width: 100%;
            border-radius: 8px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.2);
            border: 1px solid #e0e0e0;
        }
        .screenshot-tip {
            color: #666;
            margin-bottom: 10px;
            font-size: 14px;
        }
        
        .spinner {
            border: 3px solid #f3f3f3;
            border-top: 3px solid #667eea;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 1s linear infinite;
            margin: 20px auto;
        }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        
        .progress { display: none; text-align: center; padding: 20px; }
        .status-text { color: #666; margin-top: 10px; }
        
        .timer { color: #999; font-size: 12px; margin-top: 5px; }
    </style>
</head>
<body>
    <div class="container">
        <h1>周报助手</h1>
        <p class="subtitle">研究管理团队周报自动生成</p>
        
        <div id="msg" class="msg info" style="display:none;"></div>
        
        <form id="form">
            <label>WPS表格地址</label>
            <input type="text" id="wps_url" value="WPS_URL" placeholder="WPS在线表格地址">
            
            <label>日期范围（选填，如：0303-0307）</label>
            <input type="text" id="date" placeholder="留空自动从Sheet名称识别">
            
            <button type="submit" class="btn" id="btn">生成周报</button>
        </form>
        
        <div class="progress" id="progress">
            <div class="spinner"></div>
            <div id="status" class="status-text">正在处理...</div>
            <div id="timer" class="timer"></div>
        </div>
        
        <div class="screenshot-container" id="screenshot-container">
            <p class="screenshot-tip">📱 请扫描下方二维码登录微信</p>
            <img id="screenshot" src="" alt="登录二维码">
            <p class="timer" id="wait_tip">等待扫码中...</p>
        </div>
    </div>
    
    <script>
        let pollInterval;
        let startTime;
        
        document.getElementById('form').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const wpsUrl = document.getElementById('wps_url').value.trim();
            const date = document.getElementById('date').value.trim();
            const msgDiv = document.getElementById('msg');
            const progressDiv = document.getElementById('progress');
            const statusDiv = document.getElementById('status');
            const timerDiv = document.getElementById('timer');
            const screenshotContainer = document.getElementById('screenshot-container');
            const screenshotImg = document.getElementById('screenshot');
            const waitTip = document.getElementById('wait_tip');
            const btn = document.getElementById('btn');
            
            btn.disabled = true;
            msgDiv.style.display = 'none';
            progressDiv.style.display = 'block';
            screenshotContainer.style.display = 'none';
            startTime = Date.now();
            
            try {
                // 开始生成
                statusDiv.textContent = '正在启动浏览器...';
                const startRes = await fetch('/api/start', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ wps_url: wpsUrl, date: date })
                });
                const startData = await startRes.json();
                const taskId = startData.task_id;
                
                // 轮询状态
                let done = false;
                while (!done) {
                    await new Promise(r => setTimeout(r, 1500));
                    
                    const checkRes = await fetch('/api/status/' + taskId);
                    const data = await checkRes.json();
                    
                    // 更新状态
                    statusDiv.textContent = data.message || '处理中...';
                    
                    // 计时器
                    const elapsed = Math.floor((Date.now() - startTime) / 1000);
                    timerDiv.textContent = '已用时: ' + elapsed + ' 秒';
                    
                    if (data.status === 'waiting_login') {
                        // 显示二维码
                        progressDiv.style.display = 'none';
                        screenshotContainer.style.display = 'block';
                        
                        if (data.screenshot) {
                            screenshotImg.src = 'data:image/jpeg;base64,' + data.screenshot;
                        }
                        if (data.wait_time) {
                            waitTip.textContent = '等待扫码中... ' + data.wait_time + ' 秒';
                        }
                        
                    } else if (data.status === 'done') {
                        done = true;
                        screenshotContainer.style.display = 'none';
                        progressDiv.style.display = 'none';
                        
                        // 下载文档
                        window.location.href = '/api/download/' + taskId;
                        
                        msgDiv.className = 'msg success';
                        msgDiv.textContent = '周报已生成，正在下载...';
                        msgDiv.style.display = 'block';
                        
                    } else if (data.status === 'error') {
                        throw new Error(data.message);
                    }
                }
                
            } catch (err) {
                msgDiv.className = 'msg error';
                msgDiv.textContent = '错误: ' + err.message;
                msgDiv.style.display = 'block';
                progressDiv.style.display = 'none';
                screenshotContainer.style.display = 'none';
            } finally {
                btn.disabled = false;
            }
        });
    </script>
</body>
</html>
""".replace("WPS_URL", DEFAULT_WPS_URL)


@app.get("/", response_class=HTMLResponse)
async def index():
    return HTML_PAGE


class StartRequest(BaseModel):
    wps_url: str = DEFAULT_WPS_URL
    date: str = ""


@app.post("/api/start")
async def start_task(req: StartRequest):
    """开始任务"""
    task_id = datetime.now().strftime("%Y%m%d%H%M%S") + str(id(req))
    tasks[task_id] = {
        "status": "starting",
        "message": "正在初始化..."
    }
    
    # 启动后台任务
    asyncio.create_task(run_browser_task(task_id, req.wps_url, req.date))
    
    return {"task_id": task_id}


@app.get("/api/status/{task_id}")
async def get_status(task_id: str):
    """获取任务状态"""
    task = tasks.get(task_id, {"status": "unknown", "message": "任务不存在"})
    # 不返回 document 字段（太大）
    result = {k: v for k, v in task.items() if k != "document"}
    return result


@app.get("/api/download/{task_id}")
async def download(task_id: str):
    """下载文档"""
    task = tasks.get(task_id, {})
    if task.get("document"):
        doc_b64 = task["document"]
        doc_bytes = base64.b64decode(doc_b64)
        date = task.get("date", "unknown")
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="周报({date}).docx"'}
        )
    return JSONResponse({"error": "文档不存在"}, status_code=404)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
