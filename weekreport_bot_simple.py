"""
周报自动汇总助手 - 简化版
无需安装额外浏览器，直接使用系统默认浏览器
"""

import os
import sys
import io
import webbrowser
from datetime import datetime, timedelta
from pathlib import Path

# 设置UTF-8编码输出
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("[错误] 请先安装: pip install python-docx")
    input("按回车退出...")
    exit(1)

# 配置
WPS_TABLE_URL = "https://www.kdocs.cn/l/cmvjNWclJM5P"
HISTORY_WEEKLY_PATH = Path("d:/MyProject/demo_week")
OUTPUT_PATH = Path("d:/MyProject/Agent_weekreport_Bot/output")


def read_history_weekly():
    """读取历史周报模板"""
    print("\n" + "="*50)
    print("[读取历史周报模板]")
    print("="*50)
    
    templates = []
    for file in HISTORY_WEEKLY_PATH.glob("*.docx"):
        try:
            doc = Document(str(file))
            content = []
            for para in doc.paragraphs:
                content.append(para.text)
            templates.append({
                'filename': file.name,
                'content': '\n'.join(content),
                'paragraphs': [p.text for p in doc.paragraphs],
                'tables': [[cell.text for cell in row.cells] for table in doc.tables for row in table.rows]
            })
            print(f"[OK] {file.name}")
        except Exception as e:
            print(f"[跳过] {file.name} ({e})")
    
    return templates


def parse_table_content(content):
    """解析粘贴的表格内容"""
    lines = content.strip().split('\n')
    data = []
    for line in lines:
        if line.strip():
            # 按制表符或多个空格分割
            cells = [c.strip() for c in line.split('\t') if c.strip()]
            if not cells:
                cells = [c.strip() for c in line.split('  ') if c.strip()]
            if cells:
                data.append(cells)
    return data


def generate_weekly_report(table_content, templates):
    """生成周报文档"""
    print("\n" + "="*50)
    print("[生成周报文档]")
    print("="*50)
    
    OUTPUT_PATH.mkdir(parents=True, exist_ok=True)
    
    # 计算本周日期范围
    today = datetime.now()
    monday = today - timedelta(days=today.weekday())
    friday = monday + timedelta(days=4)
    date_range = f"{monday.strftime('%m%d')}-{friday.strftime('%m%d')}"
    
    doc = Document()
    
    # 标题
    title = doc.add_heading(f'研究管理团队周报（{date_range}）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 添加日期信息
    doc.add_paragraph(f"报告周期：{monday.strftime('%Y年%m月%d日')} - {friday.strftime('%Y年%m月%d日')}")
    doc.add_paragraph()
    
    # 解析并添加表格内容
    if table_content:
        data = parse_table_content(table_content)
        
        # 尝试识别模块
        current_section = None
        section_content = []
        
        for row in data:
            if len(row) > 0:
                first_cell = row[0]
                
                # 检查是否是标题行（包含"工作"、"计划"等关键词）
                if any(kw in first_cell for kw in ['工作', '计划', '问题', '风险', '协调']):
                    # 如果有之前的内容，先写入
                    if current_section and section_content:
                        doc.add_heading(current_section, level=1)
                        for item in section_content:
                            doc.add_paragraph(item)
                        section_content = []
                    
                    current_section = first_cell
                    
                    # 如果该行还有其他内容
                    if len(row) > 1:
                        section_content.append(' | '.join(row[1:]))
                else:
                    # 普通内容行
                    section_content.append(' | '.join(row))
        
        # 写入最后一个模块
        if current_section:
            doc.add_heading(current_section, level=1)
            for item in section_content:
                if item.strip():
                    doc.add_paragraph(item)
        elif data:
            # 如果没有识别到标题，直接添加所有内容
            doc.add_heading('本周工作内容', level=1)
            for row in data:
                doc.add_paragraph(' | '.join(row))
    
    # 添加标准模块（如果内容中没有）
    doc.add_paragraph()
    doc.add_heading('下周工作计划', level=1)
    doc.add_paragraph('[待填写]')
    
    doc.add_heading('问题与风险', level=1)
    doc.add_paragraph('[待填写]')
    
    doc.add_heading('需要协调事项', level=1)
    doc.add_paragraph('[待填写]')
    
    # 保存
    output_file = OUTPUT_PATH / f"研究管理团队周报（{date_range}）.docx"
    doc.save(str(output_file))
    print(f"\n[完成] 周报已保存: {output_file}")
    
    return output_file


def main():
    print("\n" + "="*60)
    print("       周报自动汇总助手")
    print("="*60)
    
    # 1. 读取历史周报模板
    templates = read_history_weekly()
    
    # 2. 打开WPS表格
    print("\n" + "="*50)
    print("[打开WPS表格]")
    print("="*50)
    print(f"\n正在打开: {WPS_TABLE_URL}")
    print("请在浏览器中:")
    print("  1. 扫码登录（如需要）")
    print("  2. 选择最新的Sheet")
    print("  3. 全选表格内容 (Ctrl+A)")
    print("  4. 复制 (Ctrl+C)")
    print("  5. 回到此窗口粘贴\n")
    
    webbrowser.open(WPS_TABLE_URL)
    
    # 3. 等待用户粘贴内容
    print("="*50)
    print("[请粘贴表格内容]")
    print("="*50)
    print("粘贴完成后按 Ctrl+Z 然后回车（或输入 END 结束）:\n")
    
    lines = []
    try:
        while True:
            line = input()
            if line.strip() == 'END':
                break
            lines.append(line)
    except EOFError:
        pass
    
    table_content = '\n'.join(lines)
    
    if not table_content.strip():
        print("\n[提示] 未检测到内容，生成空白模板...")
    
    # 4. 生成周报
    output_file = generate_weekly_report(table_content, templates)
    
    # 5. 打开输出目录
    print(f"\n[打开输出目录]")
    os.startfile(str(OUTPUT_PATH))
    
    print("\n" + "="*60)
    print("       完成!")
    print("="*60)
    input("\n按回车退出...")


if __name__ == "__main__":
    main()
