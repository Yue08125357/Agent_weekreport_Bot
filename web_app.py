"""
周报助手 - Web版
运行: python web_app.py
访问: http://localhost:8000
"""

import io
import sys
import csv
import re
import shutil
import tempfile
from pathlib import Path
from datetime import datetime, timedelta
from io import StringIO

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, FileResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = FastAPI(title="周报助手")

# 配置 - 自动判断本地/云端环境
import os
if os.path.exists("d:/MyProject/demo_week/研究管理团队周报（0126-0130).docx"):
    TEMPLATE = Path("d:/MyProject/demo_week/研究管理团队周报（0126-0130).docx")
else:
    TEMPLATE = Path("template.docx")  # 云端使用
WPS_URL = "https://www.kdocs.cn/l/cmvjNWclJM5P"


class ReportRequest(BaseModel):
    table_content: str
    date: str


def parse_table(text):
    """解析表格"""
    reader = csv.reader(StringIO(text), delimiter='\t')
    rows = []
    for row in reader:
        rows.append([cell.strip() for cell in row])
    return rows


def set_font(cell, font='仿宋', size=11, bold=False):
    """设置字体"""
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = font
            r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            r.font.size = Pt(size)
            r.font.bold = bold


def add_borders(table):
    """加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for n in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)
    if not tbl.tblPr:
        tbl.insert(0, tblPr)


def generate_report(table_content: str, date: str) -> bytes:
    """生成周报文档"""
    data = parse_table(table_content)
    
    if not data:
        raise ValueError("表格内容为空")
    
    # 复制模板
    doc = Document(str(TEMPLATE))
    
    # 删除旧表格
    for t in list(doc.tables):
        t._tbl.getparent().remove(t._tbl)
    
    # 修改标题日期
    for para in doc.paragraphs:
        full_text = ''.join([r.text for r in para.runs])
        if '工作周报' in full_text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = f'工作周报（{date}）'
            para.runs[0].font.name = '微软雅黑'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            para.runs[0].font.size = Pt(20)
            para.runs[0].font.bold = True
            break
    
    # 找详细工单位置
    detail = None
    for para in doc.paragraphs:
        if '详细工单' in para.text:
            detail = para
            break
    
    # 创建表格
    cols = max(len(r) for r in data)
    
    # 表格1
    t1 = doc.add_table(len(data), cols)
    add_borders(t1)
    for i, row in enumerate(data):
        for j, txt in enumerate(row):
            t1.rows[i].cells[j].text = txt
            set_font(t1.rows[i].cells[j], bold=(i==0))
    
    if detail:
        detail._element.addprevious(t1._tbl)
    
    # 表格2
    data2 = data + [['其他需汇报信息：无', '', '', '']]
    t2 = doc.add_table(len(data2), cols)
    add_borders(t2)
    for i, row in enumerate(data2):
        for j, txt in enumerate(row):
            t2.rows[i].cells[j].text = txt
            set_font(t2.rows[i].cells[j], bold=(i==0))
    
    if detail:
        detail._element.addnext(t2._tbl)
    
    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


# HTML页面
HTML_PAGE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>周报助手</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Microsoft YaHei', sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
        }
        .container {
            background: white;
            border-radius: 16px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            padding: 40px;
            max-width: 800px;
            width: 100%;
        }
        h1 {
            color: #333;
            margin-bottom: 10px;
            font-size: 28px;
        }
        .subtitle {
            color: #666;
            margin-bottom: 30px;
            font-size: 14px;
        }
        .step {
            background: #f8f9fa;
            border-radius: 8px;
            padding: 15px 20px;
            margin-bottom: 20px;
            border-left: 4px solid #667eea;
        }
        .step-title {
            font-weight: 600;
            color: #333;
            margin-bottom: 8px;
        }
        .step-content {
            color: #555;
            font-size: 14px;
        }
        .link {
            color: #667eea;
            word-break: break-all;
        }
        label {
            display: block;
            font-weight: 600;
            color: #333;
            margin-bottom: 8px;
        }
        input[type="text"] {
            width: 100%;
            padding: 12px 16px;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            font-size: 16px;
            transition: border-color 0.3s;
            margin-bottom: 20px;
        }
        input[type="text"]:focus {
            outline: none;
            border-color: #667eea;
        }
        textarea {
            width: 100%;
            padding: 12px 16px;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            font-size: 14px;
            font-family: 'Consolas', 'Monaco', monospace;
            resize: vertical;
            transition: border-color 0.3s;
            margin-bottom: 20px;
        }
        textarea:focus {
            outline: none;
            border-color: #667eea;
        }
        .btn {
            width: 100%;
            padding: 14px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 8px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: transform 0.2s, box-shadow 0.2s;
        }
        .btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 25px rgba(102, 126, 234, 0.4);
        }
        .btn:active {
            transform: translateY(0);
        }
        .btn:disabled {
            background: #ccc;
            cursor: not-allowed;
            transform: none;
        }
        .error {
            background: #fee;
            color: #c00;
            padding: 12px;
            border-radius: 8px;
            margin-bottom: 20px;
        }
        .success {
            background: #efe;
            color: #060;
            padding: 12px;
            border-radius: 8px;
            margin-bottom: 20px;
        }
        .loading {
            display: none;
            text-align: center;
            padding: 20px;
        }
        .spinner {
            border: 3px solid #f3f3f3;
            border-top: 3px solid #667eea;
            border-radius: 50%;
            width: 30px;
            height: 30px;
            animation: spin 1s linear infinite;
            margin: 0 auto 10px;
        }
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 周报助手</h1>
        <p class="subtitle">研究管理团队周报自动生成工具</p>
        
        <div class="step">
            <div class="step-title">步骤1: 打开WPS表格</div>
            <div class="step-content">
                点击链接：<a href="WPS_URL" target="_blank" class="link">WPS_URL</a><br>
                扫码登录后，选择最新的Sheet
            </div>
        </div>
        
        <div class="step">
            <div class="step-title">步骤2: 复制表格内容</div>
            <div class="step-content">
                在WPS表格中按 <strong>Ctrl+A</strong> 全选，再按 <strong>Ctrl+C</strong> 复制
            </div>
        </div>
        
        <div class="step">
            <div class="step-title">步骤3: 粘贴并生成</div>
            <div class="step-content">
                将复制的内容粘贴到下方文本框，输入日期后点击生成
            </div>
        </div>
        
        <div id="message"></div>
        
        <form id="reportForm">
            <label for="date">日期（如：0224-0228）</label>
            <input type="text" id="date" name="date" placeholder="输入日期，如 0224-0228">
            
            <label for="content">表格内容（粘贴复制的表格）</label>
            <textarea id="content" name="content" rows="10" placeholder="在此粘贴从WPS表格复制的内容..."></textarea>
            
            <button type="submit" class="btn" id="submitBtn">生成周报</button>
        </form>
        
        <div class="loading" id="loading">
            <div class="spinner"></div>
            <div>正在生成周报...</div>
        </div>
    </div>
    
    <script>
        // 设置默认日期（本周）
        function getDefaultDate() {
            const now = new Date();
            const day = now.getDay();
            const monday = new Date(now);
            monday.setDate(now.getDate() - (day === 0 ? 6 : day - 1));
            const friday = new Date(monday);
            friday.setDate(monday.getDate() + 4);
            
            const pad = n => n.toString().padStart(2, '0');
            return pad(monday.getMonth() + 1) + pad(monday.getDate()) + '-' + 
                   pad(friday.getMonth() + 1) + pad(friday.getDate());
        }
        
        document.getElementById('date').value = getDefaultDate();
        
        document.getElementById('reportForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const date = document.getElementById('date').value.trim();
            const content = document.getElementById('content').value.trim();
            const messageDiv = document.getElementById('message');
            const loadingDiv = document.getElementById('loading');
            const submitBtn = document.getElementById('submitBtn');
            
            if (!date) {
                messageDiv.innerHTML = '<div class="error">请输入日期</div>';
                return;
            }
            
            if (!content) {
                messageDiv.innerHTML = '<div class="error">请粘贴表格内容</div>';
                return;
            }
            
            messageDiv.innerHTML = '';
            loadingDiv.style.display = 'block';
            submitBtn.disabled = true;
            
            try {
                const response = await fetch('/api/generate', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        table_content: content,
                        date: date
                    })
                });
                
                if (!response.ok) {
                    const error = await response.json();
                    throw new Error(error.detail || '生成失败');
                }
                
                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = `研究管理团队周报（${date}）.docx`;
                document.body.appendChild(a);
                a.click();
                window.URL.revokeObjectURL(url);
                a.remove();
                
                messageDiv.innerHTML = '<div class="success">✅ 周报已生成并开始下载！</div>';
                
            } catch (err) {
                messageDiv.innerHTML = `<div class="error">❌ ${err.message}</div>`;
            } finally {
                loadingDiv.style.display = 'none';
                submitBtn.disabled = false;
            }
        });
    </script>
</body>
</html>
""".replace("WPS_URL", WPS_URL)


@app.get("/", response_class=HTMLResponse)
async def index():
    """首页"""
    return HTML_PAGE


@app.post("/api/generate")
async def generate(req: ReportRequest):
    """生成周报"""
    try:
        doc_bytes = generate_report(req.table_content, req.date)
        
        filename = f"研究管理团队周报（{req.date}）.docx"
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    import uvicorn
    print("\n" + "="*50)
    print("  周报助手 Web版")
    print("="*50)
    print(f"\n  访问地址: http://localhost:8000")
    print(f"  WPS表格: {WPS_URL}")
    print("\n  按 Ctrl+C 停止服务\n")
    uvicorn.run(app, host="0.0.0.0", port=8000)
