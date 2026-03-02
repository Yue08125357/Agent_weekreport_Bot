"""
周报助手 - 全自动版
"""
import asyncio
import sys
import io
import os
import re
import shutil
import csv
from pathlib import Path
from datetime import datetime, timedelta
from io import StringIO

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import win32clipboard

WPS_URL = "https://www.kdocs.cn/l/cmvjNWclJM5P"
TEMPLATE = Path("d:/MyProject/demo_week/研究管理团队周报（0126-0130).docx")
OUTPUT = Path("d:/MyProject/Agent_weekreport_Bot/output")
BROWSER_DATA = Path("d:/MyProject/Agent_weekreport_Bot/browser_data")


def get_clipboard():
    """读剪贴板"""
    win32clipboard.OpenClipboard()
    text = win32clipboard.GetClipboardData(win32clipboard.CF_UNICODETEXT)
    win32clipboard.CloseClipboard()
    return text


def parse_table(text):
    """解析表格"""
    reader = csv.reader(StringIO(text), delimiter='\t')
    rows = []
    for row in reader:
        rows.append([cell.strip() for cell in row])
    return rows


def set_font(cell, font='仿宋', size=11, bold=False):
    """设置字体"""
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = font
            r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            r.font.size = Pt(size)
            r.font.bold = bold


def add_borders(table):
    """加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for n in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)
    if not tbl.tblPr:
        tbl.insert(0, tblPr)


async def main():
    print("\n===== 周报助手 =====\n")
    
    # 1. 打开浏览器
    p = await async_playwright().start()
    BROWSER_DATA.mkdir(exist_ok=True)
    ctx = await p.chromium.launch_persistent_context(
        str(BROWSER_DATA), headless=False, channel="chrome",
        viewport={'width': 1400, 'height': 900}
    )
    page = ctx.pages[0] if ctx.pages else await ctx.new_page()
    
    print("打开WPS表格...")
    await page.goto(WPS_URL)
    
    # 2. 等待登录
    print("\n等待登录...")
    print("请在浏览器中扫码登录微信")
    print("登录成功后会自动继续...\n")
    
    # 等待表格加载（检测URL变化或表格元素）
    for i in range(90):
        await asyncio.sleep(2)
        url = page.url
        if 'passport' not in url and 'singlesign' not in url:
            # 检查是否有表格
            content = await page.content()
            if 'sheet' in content.lower() or 'editor' in content.lower():
                print("登录成功！")
                break
        if i % 5 == 0:
            print(f"  等待中... {i*2}秒")
    else:
        print("超时，请手动操作")
        input("准备好后按回车继续...")
    
    await asyncio.sleep(2)
    
    # 3. 获取最新Sheet名称
    print("\n获取Sheet名称...")
    sheet_name = ""
    
    # 尝试多种选择器
    for selector in ['[class*="sheet-tab"]', '.sheet-tab', '[role="tab"]']:
        try:
            tabs = await page.query_selector_all(selector)
            if tabs:
                last_tab = tabs[-1]
                sheet_name = await last_tab.inner_text()
                sheet_name = sheet_name.strip()
                # 点击切换
                await last_tab.click()
                await asyncio.sleep(2)
                print(f"Sheet名称: {sheet_name}")
                break
        except:
            continue
    
    # 提取日期
    date_match = re.search(r'(\d{4}[-—]\d{4})', sheet_name)
    if date_match:
        date = date_match.group(1)
    else:
        print(f"\n自动识别日期: {sheet_name}")
        date = input("输入日期（如0126-0130，回车使用识别结果）: ").strip()
        if not date:
            date = sheet_name
    
    print(f"日期: {date}")
    
    # 4. 自动全选并复制
    print("\n自动复制表格...")
    
    # 点击表格区域确保焦点
    try:
        await page.click('[class*="canvas"], [class*="editor"], [class*="sheet"]')
    except:
        pass
    
    await asyncio.sleep(0.5)
    
    # Ctrl+A 全选
    await page.keyboard.press('Control+A')
    await asyncio.sleep(0.5)
    
    # Ctrl+C 复制
    await page.keyboard.press('Control+C')
    await asyncio.sleep(1)
    
    print("已复制到剪贴板")
    
    # 5. 读剪贴板
    print("\n读取剪贴板...")
    text = get_clipboard()
    print(f"内容长度: {len(text)} 字符")
    
    data = parse_table(text)
    print(f"解析: {len(data)} 行 x {max(len(r) for r in data) if data else 0} 列")
    
    # 6. 生成文档
    print("\n生成文档...")
    OUTPUT.mkdir(exist_ok=True)
    out_file = OUTPUT / f"研究管理团队周报（{date}）.docx"
    shutil.copy(TEMPLATE, out_file)
    
    doc = Document(out_file)
    
    # 删除旧表格
    for t in list(doc.tables):
        t._tbl.getparent().remove(t._tbl)
    
    # 修改标题日期 - 遍历所有段落
    print("修改标题日期...")
    for para in doc.paragraphs:
        full_text = ''.join([r.text for r in para.runs])
        if '工作周报' in full_text:
            # 清空并重写
            for run in para.runs:
                run.text = ''
            # 在第一个run写入新标题
            para.runs[0].text = f'工作周报（{date}）'
            # 保持原格式
            para.runs[0].font.name = '微软雅黑'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            para.runs[0].font.size = Pt(20)
            para.runs[0].font.bold = True
            print(f"  标题: 工作周报（{date}）")
            break
    
    # 找详细工单位置
    detail = None
    for para in doc.paragraphs:
        if '详细工单' in para.text:
            detail = para
            break
    
    # 创建表格
    cols = max(len(r) for r in data)
    
    # 表格1
    t1 = doc.add_table(len(data), cols)
    add_borders(t1)
    for i, row in enumerate(data):
        for j, txt in enumerate(row):
            t1.rows[i].cells[j].text = txt
            set_font(t1.rows[i].cells[j], bold=(i==0))
    
    if detail:
        detail._element.addprevious(t1._tbl)
    
    # 表格2
    data2 = data + [['其他需汇报信息：无', '', '', '']]
    t2 = doc.add_table(len(data2), cols)
    add_borders(t2)
    for i, row in enumerate(data2):
        for j, txt in enumerate(row):
            t2.rows[i].cells[j].text = txt
            set_font(t2.rows[i].cells[j], bold=(i==0))
    
    if detail:
        detail._element.addnext(t2._tbl)
    
    doc.save(out_file)
    print(f"已保存: {out_file}")
    
    await ctx.close()
    await p.stop()
    
    os.startfile(out_file)
    print("\n完成！")


if __name__ == "__main__":
    asyncio.run(main())
