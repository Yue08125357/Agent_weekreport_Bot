"""
周报助手 - 内网穿透全自动版
流程：用户输入日期 → 自动打开WPS → 自动登录等待 → 自动复制表格 → 生成周报 → 下载
"""

import asyncio
import sys
import io
import os
import re
import csv
import shutil
from pathlib import Path
from datetime import datetime, timedelta
from io import StringIO
from typing import Optional

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, Response
from pydantic import BaseModel

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配置
TEMPLATE = Path("d:/MyProject/demo_week/研究管理团队周报（0126-0130).docx")
OUTPUT_DIR = Path("d:/MyProject/Agent_weekreport_Bot/output")
BROWSER_DATA = Path("d:/MyProject/Agent_weekreport_Bot/browser_data")
DEFAULT_WPS_URL = "https://www.kdocs.cn/l/cmvjNWclJM5P"

app = FastAPI(title="周报助手")

# 全局变量
playwright_instance = None
browser_context = None
browser_page = None


class ReportRequest(BaseModel):
    wps_url: str = DEFAULT_WPS_URL
    date: str


def parse_table(text):
    """解析表格"""
    reader = csv.reader(StringIO(text), delimiter='\t')
    rows = []
    for row in reader:
        rows.append([cell.strip() for cell in row])
    return rows


def set_font(cell, font='仿宋', size=11, bold=False):
    """设置字体"""
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = font
            r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            r.font.size = Pt(size)
            r.font.bold = bold


def add_borders(table):
    """加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for n in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)
    if not tbl.tblPr:
        tbl.insert(0, tblPr)


def generate_report(table_content: str, date: str) -> bytes:
    """生成周报文档"""
    data = parse_table(table_content)
    if not data:
        raise ValueError("表格内容为空")
    
    doc = Document(str(TEMPLATE))
    
    # 删除旧表格
    for t in list(doc.tables):
        t._tbl.getparent().remove(t._tbl)
    
    # 修改标题日期
    for para in doc.paragraphs:
        full_text = ''.join([r.text for r in para.runs])
        if '工作周报' in full_text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = f'工作周报（{date}）'
            para.runs[0].font.name = '微软雅黑'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            para.runs[0].font.size = Pt(20)
            para.runs[0].font.bold = True
            break
    
    # 找详细工单位置
    detail = None
    for para in doc.paragraphs:
        if '详细工单' in para.text:
            detail = para
            break
    
    cols = max(len(r) for r in data)
    
    # 表格1
    t1 = doc.add_table(len(data), cols)
    add_borders(t1)
    for i, row in enumerate(data):
        for j, txt in enumerate(row):
            t1.rows[i].cells[j].text = txt
            set_font(t1.rows[i].cells[j], bold=(i==0))
    
    if detail:
        detail._element.addprevious(t1._tbl)
    
    # 表格2
    data2 = data + [['其他需汇报信息：无', '', '', '']]
    t2 = doc.add_table(len(data2), cols)
    add_borders(t2)
    for i, row in enumerate(data2):
        for j, txt in enumerate(row):
            t2.rows[i].cells[j].text = txt
            set_font(t2.rows[i].cells[j], bold=(i==0))
    
    if detail:
        detail._element.addnext(t2._tbl)
    
    # 保存
    OUTPUT_DIR.mkdir(exist_ok=True)
    output_file = OUTPUT_DIR / f"研究管理团队周报（{date}）.docx"
    doc.save(str(output_file))
    
    return output_file


async def get_browser_page():
    """获取或创建浏览器页面"""
    global playwright_instance, browser_context, browser_page
    
    from playwright.async_api import async_playwright
    
    if not playwright_instance:
        playwright_instance = await async_playwright().start()
        BROWSER_DATA.mkdir(exist_ok=True)
        browser_context = await playwright_instance.chromium.launch_persistent_context(
            str(BROWSER_DATA),
            headless=False,
            channel="chrome",
            args=['--disable-blink-features=AutomationControlled'],
            ignore_default_args=['--enable-automation'],
            viewport={'width': 1400, 'height': 900}
        )
    
    if browser_context.pages:
        browser_page = browser_context.pages[0]
    else:
        browser_page = await browser_context.new_page()
    
    return browser_page


async def auto_copy_table(wps_url: str) -> tuple:
    """自动打开WPS并复制表格，返回(表格内容, Sheet名称)"""
    page = await get_browser_page()
    
    # 打开WPS
    await page.goto(wps_url)
    
    # 等待登录（检测表格加载）
    for i in range(90):
        await asyncio.sleep(2)
        url = page.url
        if 'passport' not in url and 'singlesign' not in url:
            content = await page.content()
            if 'sheet' in content.lower() or 'editor' in content.lower():
                break
    
    await asyncio.sleep(2)
    
    # 获取最新Sheet名称
    sheet_name = ""
    for selector in ['[class*="sheet-tab"]', '.sheet-tab', '[role="tab"]']:
        try:
            tabs = await page.query_selector_all(selector)
            if tabs:
                last_tab = tabs[-1]
                sheet_name = await last_tab.inner_text()
                sheet_name = sheet_name.strip()
                await last_tab.click()
                await asyncio.sleep(2)
                break
        except:
            continue
    
    # 点击表格区域
    try:
        await page.click('[class*="canvas"], [class*="editor"], [class*="sheet"]')
    except:
        pass
    
    await asyncio.sleep(0.5)
    
    # Ctrl+A 全选
    await page.keyboard.press('Control+A')
    await asyncio.sleep(0.5)
    
    # Ctrl+C 复制
    await page.keyboard.press('Control+C')
    await asyncio.sleep(1)
    
    # 读取剪贴板
    import win32clipboard
    win32clipboard.OpenClipboard()
    text = win32clipboard.GetClipboardData(win32clipboard.CF_UNICODETEXT)
    win32clipboard.CloseClipboard()
    
    return text, sheet_name


# HTML页面
HTML_PAGE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>周报助手</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Microsoft YaHei', sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
        }
        .container {
            background: white;
            border-radius: 16px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            padding: 40px;
            max-width: 600px;
            width: 100%;
        }
        h1 {
            color: #333;
            margin-bottom: 10px;
            font-size: 28px;
            text-align: center;
        }
        .subtitle {
            color: #666;
            margin-bottom: 30px;
            font-size: 14px;
            text-align: center;
        }
        label {
            display: block;
            font-weight: 600;
            color: #333;
            margin-bottom: 8px;
        }
        input[type="text"] {
            width: 100%;
            padding: 12px 16px;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            font-size: 16px;
            margin-bottom: 20px;
        }
        input[type="text"]:focus {
            outline: none;
            border-color: #667eea;
        }
        .btn {
            width: 100%;
            padding: 14px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 8px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: transform 0.2s;
        }
        .btn:hover {
            transform: translateY(-2px);
        }
        .btn:disabled {
            background: #ccc;
            cursor: not-allowed;
            transform: none;
        }
        .msg {
            padding: 12px;
            border-radius: 8px;
            margin-bottom: 20px;
            display: none;
        }
        .msg.error { background: #fee; color: #c00; display: block; }
        .msg.success { background: #efe; color: #060; display: block; }
        .msg.info { background: #e3f2fd; color: #1565c0; display: block; }
        .loading {
            text-align: center;
            padding: 20px;
            display: none;
        }
        .spinner {
            border: 3px solid #f3f3f3;
            border-top: 3px solid #667eea;
            border-radius: 50%;
            width: 30px;
            height: 30px;
            animation: spin 1s linear infinite;
            margin: 0 auto 10px;
        }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        .note {
            background: #fff3e0;
            padding: 12px;
            border-radius: 8px;
            margin-bottom: 20px;
            font-size: 13px;
            color: #e65100;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>周报助手</h1>
        <p class="subtitle">研究管理团队周报自动生成</p>
        
        <div class="note">
            点击生成后，请在弹出的浏览器窗口中扫码登录微信
        </div>
        
        <div id="msg" class="msg"></div>
        
        <form id="form">
            <label>WPS表格地址</label>
            <input type="text" id="wps_url" value="WPS_URL" placeholder="WPS在线表格地址">
            
            <label>日期范围（如：0224-0228）</label>
            <input type="text" id="date" placeholder="留空自动识别">
            
            <button type="submit" class="btn" id="btn">生成周报</button>
        </form>
        
        <div class="loading" id="loading">
            <div class="spinner"></div>
            <div id="status">正在处理...</div>
        </div>
    </div>
    
    <script>
        function getDefaultDate() {
            const now = new Date();
            const day = now.getDay();
            const monday = new Date(now);
            monday.setDate(now.getDate() - (day === 0 ? 6 : day - 1));
            const friday = new Date(monday);
            friday.setDate(monday.getDate() + 4);
            const pad = n => n.toString().padStart(2, '0');
            return pad(monday.getMonth() + 1) + pad(monday.getDate()) + '-' + 
                   pad(friday.getMonth() + 1) + pad(friday.getDate());
        }
        
        document.getElementById('date').placeholder = '留空使用本周: ' + getDefaultDate();
        
        document.getElementById('form').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const wpsUrl = document.getElementById('wps_url').value.trim();
            const date = document.getElementById('date').value.trim();
            const msgDiv = document.getElementById('msg');
            const loadingDiv = document.getElementById('loading');
            const statusDiv = document.getElementById('status');
            const btn = document.getElementById('btn');
            
            msgDiv.className = 'msg';
            msgDiv.textContent = '';
            loadingDiv.style.display = 'block';
            btn.disabled = true;
            
            try {
                statusDiv.textContent = '正在打开WPS表格...';
                
                const response = await fetch('/api/generate', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ wps_url: wpsUrl, date: date })
                });
                
                const data = await response.json();
                
                if (data.error) {
                    throw new Error(data.error);
                }
                
                if (data.status === 'waiting_login') {
                    msgDiv.className = 'msg info';
                    msgDiv.textContent = '请在弹出的浏览器窗口中扫码登录微信...';
                    statusDiv.textContent = '等待登录...';
                    
                    // 轮询检查状态
                    let result = null;
                    for (let i = 0; i < 60; i++) {
                        await new Promise(r => setTimeout(r, 3000));
                        statusDiv.textContent = `等待登录... (${i * 3}秒)`;
                        
                        const checkRes = await fetch('/api/check?id=' + data.task_id);
                        const checkData = await checkRes.json();
                        
                        if (checkData.status === 'done') {
                            result = checkData;
                            break;
                        } else if (checkData.status === 'error') {
                            throw new Error(checkData.error);
                        }
                    }
                    
                    if (!result) {
                        throw new Error('超时，请重试');
                    }
                    
                    // 下载文件
                    window.location.href = '/api/download?id=' + data.task_id;
                    msgDiv.className = 'msg success';
                    msgDiv.textContent = '周报已生成，正在下载...';
                    
                } else if (data.download_url) {
                    window.location.href = data.download_url;
                    msgDiv.className = 'msg success';
                    msgDiv.textContent = '周报已生成！';
                }
                
            } catch (err) {
                msgDiv.className = 'msg error';
                msgDiv.textContent = '错误: ' + err.message;
            } finally {
                loadingDiv.style.display = 'none';
                btn.disabled = false;
            }
        });
    </script>
</body>
</html>
""".replace("WPS_URL", DEFAULT_WPS_URL)


# 任务存储
tasks = {}


@app.get("/", response_class=HTMLResponse)
async def index():
    return HTML_PAGE


@app.post("/api/generate")
async def generate(req: ReportRequest):
    """生成周报"""
    task_id = datetime.now().strftime("%Y%m%d%H%M%S")
    tasks[task_id] = {"status": "waiting_login"}
    
    # 启动后台任务
    asyncio.create_task(process_report(task_id, req.wps_url, req.date))
    
    return {"status": "waiting_login", "task_id": task_id}


@app.get("/api/check")
async def check_status(task_id: str):
    """检查任务状态"""
    return tasks.get(task_id, {"status": "unknown"})


@app.get("/api/download")
async def download(task_id: str):
    """下载文件"""
    task = tasks.get(task_id, {})
    if task.get("file"):
        file_path = Path(task["file"])
        if file_path.exists():
            return FileResponse(
                file_path,
                filename=file_path.name,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    return {"error": "文件不存在"}


async def process_report(task_id: str, wps_url: str, date: str):
    """后台处理"""
    try:
        # 自动复制表格
        table_content, sheet_name = await auto_copy_table(wps_url)
        
        # 确定日期
        if not date:
            match = re.search(r'(\d{4}[-—]\d{4})', sheet_name)
            date = match.group(1) if match else sheet_name
        
        # 生成报告
        output_file = generate_report(table_content, date)
        
        tasks[task_id] = {
            "status": "done",
            "file": str(output_file),
            "date": date
        }
        
    except Exception as e:
        tasks[task_id] = {"status": "error", "error": str(e)}


if __name__ == "__main__":
    import uvicorn
    print("\n" + "="*50)
    print("  周报助手 - 内网穿透版")
    print("="*50)
    print("\n  访问地址: http://localhost:8000")
    print("\n  启动 ngrok 后可获得公网地址")
    print("\n  按 Ctrl+C 停止\n")
    uvicorn.run(app, host="0.0.0.0", port=8000)
